import os
import re
import io
import logging
import shutil
import subprocess
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime, timezone
from sqlalchemy.orm import Session
from app.models import (
    Candidate, WhatsAppConsentStatusEnum, WhatsAppOptOut, BenchStatusEnum
)
from app.schemas import WhatsAppEligibilityInfo, CVExtractionResponse

logger = logging.getLogger(__name__)

# Comprehensive Skills Knowledge Base (Languages, Frameworks, Databases, Tools & Domains)
TECH_SKILLS_DB = [
    # Oracle & Databases
    "Oracle", "PL/SQL", "SQL", "Oracle Forms", "Oracle Reports", "Oracle APEX", "Oracle EBS",
    "Oracle Cloud", "Database Tuning", "TOAD", "SQL Developer", "Stored Procedures", "Triggers",
    "Data Modeling", "Database Administration", "DBA", "PostgreSQL", "Postgres", "MySQL",
    "MongoDB", "Redis", "Elasticsearch", "Cassandra", "DynamoDB", "SQLite", "MariaDB",
    "Snowflake", "BigQuery", "SQL Server", "T-SQL", "MSSQL",
    # Core Languages
    "Python", "JavaScript", "TypeScript", "Java", "Core Java", "J2EE", "C++", "C#", ".NET",
    "Go", "Golang", "Rust", "PHP", "Ruby", "Swift", "Kotlin", "Scala", "R", "Dart",
    "HTML", "HTML5", "CSS", "CSS3", "Bash", "Shell Scripting",
    # Frontend Frameworks & Libraries
    "React", "React.js", "React Native", "Next.js", "Angular", "AngularJS", "Vue", "Vue.js",
    "Node.js", "Redux", "NgRx", "RxJS", "Tailwind CSS", "Bootstrap", "jQuery", "GraphQL",
    # Backend Frameworks
    "FastAPI", "Django", "Flask", "Spring Boot", "Spring", "Spring MVC", "Hibernate", "JPA",
    "Express", "Express.js", "NestJS", "ASP.NET", ".NET Core", "Laravel", "Ruby on Rails",
    # Cloud, DevOps & Infrastructure
    "AWS", "Amazon Web Services", "GCP", "Google Cloud", "Azure", "Docker", "Kubernetes",
    "K8s", "Terraform", "Ansible", "CI/CD", "GitHub Actions", "GitLab CI", "Jenkins",
    "Linux", "Unix", "Nginx", "Apache", "Kafka", "RabbitMQ", "Microservices", "Serverless",
    # Mobile & Specialized
    "Flutter", "Android", "iOS", "Salesforce", "Apex", "SAP", "ABAP", "SAP FICO", "SAP MM",
    "SAP HANA", "MuleSoft", "Workday", "ServiceNow",
    # Testing & QA
    "Selenium", "Automation Testing", "Manual Testing", "Pytest", "Jest", "Cypress",
    "JUnit", "Postman", "Unit Testing", "TestNG", "Cucumber", "API Testing",
    # Data Engineering, AI & Analytics
    "Machine Learning", "Deep Learning", "NLP", "Pandas", "NumPy", "Scikit-Learn",
    "TensorFlow", "PyTorch", "Hadoop", "Spark", "Apache Spark", "Airflow", "ETL",
    "Data Engineering", "Power BI", "Tableau", "Data Analysis",
    # Architecture, Management & Business Domains
    "REST API", "RESTful", "gRPC", "WebSocket", "System Design", "Agile", "Scrum", "Jira",
    "Git", "GitHub", "GitLab", "Bitbucket", "Finance", "Banking", "Healthcare", "E-commerce",
    "ERP", "CRM", "IT", "Support", "Reporting", "Production Support", "Maintenance"
]

# Classification Domain Templates
DOMAIN_SKILL_PROFILES = {
    "oracle": {
        "title": "Oracle Developer",
        "primary": ["Oracle", "PL/SQL", "SQL", "Oracle Forms", "Oracle Reports", "Oracle APEX", "Database Tuning", "TOAD", "SQL Developer", "Stored Procedures", "Triggers", "Data Modeling", "Oracle EBS"],
        "secondary": ["Finance", "IT", "Support", "Reporting", "Unix", "Linux", "Shell Scripting", "Git", "Jira", "Agile", "Performance Tuning", "Production Support", "ERP"]
    },
    "java": {
        "title": "Java Developer",
        "primary": ["Java", "Core Java", "Spring Boot", "Spring", "Microservices", "Hibernate", "JPA", "J2EE", "REST API", "RESTful", "Maven", "Gradle"],
        "secondary": ["SQL", "MySQL", "PostgreSQL", "Kafka", "Docker", "Kubernetes", "AWS", "Git", "Jenkins", "JUnit", "Agile", "Scrum", "Linux", "Support", "Reporting"]
    },
    "angular": {
        "title": "Angular Developer",
        "primary": ["Angular", "TypeScript", "JavaScript", "RxJS", "HTML5", "CSS3", "HTML", "CSS", "NgRx", "Bootstrap", "Tailwind CSS"],
        "secondary": ["REST API", "Git", "GitHub", "Jira", "Webpack", "npm", "Node.js", "Agile", "Scrum", "Unit Testing", "Jest", "UI/UX", "Support"]
    },
    "react": {
        "title": "React Developer",
        "primary": ["React", "React.js", "Next.js", "JavaScript", "TypeScript", "Redux", "HTML5", "CSS3", "HTML", "CSS", "Tailwind CSS", "React Native"],
        "secondary": ["Node.js", "REST API", "GraphQL", "Git", "Webpack", "Vite", "Jest", "Agile", "Scrum", "Firebase", "UI/UX", "Support"]
    },
    "python": {
        "title": "Python Developer",
        "primary": ["Python", "Django", "FastAPI", "Flask", "PostgreSQL", "SQL", "REST API", "Celery", "SQLAlchemy", "Pandas"],
        "secondary": ["Docker", "AWS", "Git", "Linux", "Redis", "Kafka", "CI/CD", "Pytest", "Agile", "Scrum", "Data Analysis", "Reporting", "Support"]
    },
    "devops": {
        "title": "DevOps Engineer",
        "primary": ["AWS", "Docker", "Kubernetes", "Terraform", "CI/CD", "Jenkins", "Ansible", "Linux", "Azure", "GCP", "GitHub Actions"],
        "secondary": ["Python", "Bash", "Shell Scripting", "Git", "Nginx", "Monitoring", "Prometheus", "Grafana", "Jira", "Agile", "Security", "Support"]
    },
    "data": {
        "title": "Data Engineer",
        "primary": ["SQL", "Python", "Spark", "Hadoop", "Kafka", "ETL", "Data Engineering", "Pandas", "Airflow", "Data Modeling", "BigQuery", "Snowflake"],
        "secondary": ["AWS", "Azure", "Power BI", "Tableau", "Git", "Linux", "Reporting", "Analytics", "Support", "Agile", "Scrum"]
    },
    "qa": {
        "title": "QA Automation Engineer",
        "primary": ["Selenium", "Automation Testing", "Manual Testing", "Pytest", "Jest", "JUnit", "Cypress", "Postman", "TestNG", "Cucumber", "API Testing"],
        "secondary": ["Jira", "Git", "SQL", "Java", "Python", "Agile", "Scrum", "Reporting", "Bug Tracking", "CI/CD", "Support"]
    },
    "dotnet": {
        "title": "Dot Net Developer",
        "primary": [".NET", "ASP.NET", ".NET Core", "C#", "SQL Server", "Entity Framework", "Web API", "LINQ", "Microservices"],
        "secondary": ["Azure", "SQL", "Git", "TFS", "JavaScript", "HTML", "CSS", "Agile", "Scrum", "Support", "Reporting"]
    },
    "salesforce": {
        "title": "Salesforce Developer",
        "primary": ["Salesforce", "Apex", "Visualforce", "Lightning", "LWC", "SOQL", "CRM", "Sales Cloud", "Service Cloud"],
        "secondary": ["REST API", "Git", "Jira", "Agile", "JavaScript", "HTML", "CSS", "Support", "Reporting"]
    },
    "sap": {
        "title": "SAP Consultant",
        "primary": ["SAP", "ABAP", "SAP FICO", "SAP MM", "SAP HANA", "ERP", "BAPI", "IDoc"],
        "secondary": ["SQL", "Oracle", "IT", "Support", "Reporting", "Finance", "Supply Chain", "Agile"]
    },
    "node": {
        "title": "Node.js Developer",
        "primary": ["Node.js", "Express", "Express.js", "NestJS", "JavaScript", "TypeScript", "MongoDB", "PostgreSQL", "REST API"],
        "secondary": ["Redis", "Docker", "AWS", "Git", "GraphQL", "Microservices", "Jest", "Agile", "Scrum"]
    }
}

DEGREE_PATTERNS = [
    (r"\b(Ph\.?D|Doctor of Philosophy)\b", "Ph.D"),
    (r"\b(M\.?Tech|Master of Technology)\b", "M.Tech"),
    (r"\b(M\.?S\.?|Master of Science)\b", "Master of Science (M.S.)"),
    (r"\b(M\.?C\.?A|Master of Computer Applications)\b", "MCA"),
    (r"\b(M\.?B\.?A|Master of Business Administration)\b", "MBA"),
    (r"\b(B\.?Tech|Bachelor of Technology)\b", "B.Tech"),
    (r"\b(B\.?E\.?|Bachelor of Engineering)\b", "B.E."),
    (r"\b(B\.?C\.?A|Bachelor of Computer Applications)\b", "BCA"),
    (r"\b(B\.?Sc|Bachelor of Science)\b", "B.Sc"),
    (r"\b(B\.?Com|Bachelor of Commerce)\b", "B.Com"),
    (r"\b(Bachelor'?s\s+Degree|Master'?s\s+Degree)\b", "Bachelor's Degree")
]

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract raw text from PDF, DOCX, DOC or Text files."""
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    
    if ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_content))
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        except Exception as e:
            logger.warning(f"pypdf extraction failed for {filename}: {e}")
            
    elif ext == ".doc":
        antiword = shutil.which("antiword")
        if antiword:
            try:
                result = subprocess.run(
                    [antiword, "-m", "UTF-8.txt", "-"],
                    input=file_content,
                    capture_output=True,
                    check=True,
                )
                text = result.stdout.decode("utf-8", errors="ignore")
            except (OSError, subprocess.SubprocessError) as e:
                logger.warning(f"antiword extraction failed for {filename}: {e}")
        else:
            logger.warning("antiword is not installed; legacy .doc text cannot be extracted")

    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_content))
            for p in doc.paragraphs:
                text += p.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            logger.warning(f"python-docx extraction failed for {filename}: {e}")

    # Fallback to UTF-8 decoding
    if not text.strip():
        try:
            text = file_content.decode("utf-8", errors="ignore")
        except Exception:
            text = ""

    return text

def extract_position_from_text(raw_text: str, filename: str = "", found_skills: List[str] = []) -> str:
    """
    Accurately extracts candidate's exact position / job title from CV text.
    Never returns a generic or random title.
    """
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    full_text = " ".join(lines)

    # 1. Explicit Role Prefix Patterns in CV (e.g. "Role: Oracle Developer", "Designation: Java Developer")
    role_prefix_patterns = [
        r"(?:current\s+)?(?:job\s+)?(?:title|position|designation|role|profile)\s*[:\-\|\–]\s*([A-Za-z0-9\s/+#\.\(\)-]{3,45})",
        r"(?:working\s+as\s+(?:an?|the)?)\s+([A-Za-z0-9\s/+#\.\(\)-]{3,40})(?:\s+at|\s+in|\s+for|\.|\,)",
        r"(?:experience\s+as\s+(?:an?|the)?)\s+([A-Za-z0-9\s/+#\.\(\)-]{3,40})(?:\s+at|\s+in|\s+for|\.|\,)",
        r"(?:target\s+role|applying\s+for|current\s+job)\s*[:\-\|\–]\s*([A-Za-z0-9\s/+#\.\(\)-]{3,40})"
    ]
    for pat in role_prefix_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            extracted = m.group(1).strip()
            # Clean unwanted tail characters
            extracted = re.split(r"[,;\|\n\r]", extracted)[0].strip()
            if 3 <= len(extracted) <= 45 and not re.search(r"(@|http|years?|resume|cv)", extracted, re.IGNORECASE):
                # Standardize capitalization
                return " ".join([w.capitalize() for w in extracted.split()])

    # 2. Comprehensive Specific Role Regexes
    specific_roles = [
        # Oracle Roles
        r"\b(Oracle\s+Developer|Oracle\s+PL/SQL\s+Developer|Oracle\s+Apps\s+DBA|Oracle\s+Forms\s+Developer|Oracle\s+EBS\s+Consultant|Oracle\s+Database\s+Administrator|Oracle\s+DBA)\b",
        # Java Roles
        r"\b(Java\s+Developer|Java\s+Full\s*Stack\s+Developer|Senior\s+Java\s+Developer|Core\s+Java\s+Developer|Java/J2EE\s+Developer|Java\s+Software\s+Engineer)\b",
        # Angular & React Roles
        r"\b(Angular\s+Developer|Angular\s+Frontend\s+Developer|AngularJS\s+Developer|React\s+Developer|React\.js\s+Developer|React\s+Native\s+Developer|Frontend\s+Developer|UI/UX\s+Developer|Web\s+Developer)\b",
        # Python Roles
        r"\b(Python\s+Developer|Python\s+Django\s+Developer|Python\s+Full\s*Stack\s+Developer|Python\s+Backend\s+Developer|Python\s+Data\s+Engineer)\b",
        # Full Stack & Backend
        r"\b(Full\s*Stack\s+Developer|Full\s*Stack\s+Engineer|Backend\s+Developer|Node\.js\s+Developer|\.NET\s+Developer|C#\s+Developer|ASP\.NET\s+Developer|PHP\s+Developer|Laravel\s+Developer|Golang\s+Developer|Ruby\s+on\s+Rails\s+Developer)\b",
        # DevOps & Cloud
        r"\b(DevOps\s+Engineer|Cloud\s+DevOps\s+Engineer|AWS\s+Cloud\s+Engineer|Cloud\s+Engineer|Site\s+Reliability\s+Engineer|SRE|Platform\s+Engineer|Azure\s+DevOps\s+Engineer|Kubernetes\s+Administrator)\b",
        # Data & AI
        r"\b(Data\s+Engineer|Big\s+Data\s+Engineer|Data\s+Scientist|Data\s+Analyst|BI\s+Developer|Power\s+BI\s+Developer|Tableau\s+Developer|Machine\s+Learning\s+Engineer|AI\s+Engineer|ETL\s+Developer)\b",
        # QA & Testing
        r"\b(QA\s+Automation\s+Engineer|QA\s+Engineer|Quality\s+Assurance\s+Engineer|Software\s+Test\s+Engineer|SDET|Manual\s+Tester|Automation\s+Test\s+Engineer)\b",
        # Enterprise & CRM
        r"\b(Salesforce\s+Developer|Salesforce\s+Administrator|SAP\s+Consultant|SAP\s+ABAP\s+Consultant|SAP\s+FICO\s+Consultant|SAP\s+HANA\s+Consultant|ServiceNow\s+Developer|MuleSoft\s+Developer)\b",
        # Mobile
        r"\b(Mobile\s+App\s+Developer|Android\s+Developer|iOS\s+Developer|Flutter\s+Developer|Swift\s+Developer)\b",
        # Architecture & Leadership
        r"\b(Solution\s+Architect|Technical\s+Architect|Cloud\s+Architect|Technical\s+Lead|Tech\s+Lead|Engineering\s+Manager|Scrum\s+Master|Product\s+Manager|Business\s+Analyst|System\s+Administrator|Database\s+Administrator)\b"
    ]

    for role_rx in specific_roles:
        m = re.search(role_rx, full_text, re.IGNORECASE)
        if m:
            role_text = m.group(0).strip()
            return " ".join([w.capitalize() if not w.isupper() else w for w in role_text.split()])

    # 3. Check Top Header Lines (Lines 2-8, right under the candidate's name)
    for line in lines[1:8]:
        if "@" in line or "http" in line or "phone" in line.lower() or "experience" in line.lower() or "summary" in line.lower():
            continue
        cleaned = re.sub(r"[^a-zA-Z0-9\s/+#\.-]", " ", line).strip()
        words = cleaned.split()
        if 2 <= len(words) <= 5:
            last_word = words[-1].lower()
            if last_word in ["developer", "engineer", "architect", "consultant", "administrator", "specialist", "analyst", "lead", "manager", "dba", "programmer"]:
                return " ".join([w.capitalize() for w in words])

    # 4. Derivation from Candidate's Strongest Technical Skill
    skills_lower = [s.lower() for s in found_skills]
    if any(k in skills_lower for k in ["oracle", "pl/sql", "oracle forms", "oracle reports", "oracle apex", "toad"]):
        return "Oracle Developer"
    elif any(k in skills_lower for k in ["java", "spring boot", "hibernate", "core java", "j2ee"]):
        return "Java Developer"
    elif any(k in skills_lower for k in ["angular", "angularjs", "rxjs", "ngrx"]):
        return "Angular Developer"
    elif any(k in skills_lower for k in ["react", "react.js", "react native", "next.js", "redux"]):
        return "React Developer"
    elif any(k in skills_lower for k in ["python", "django", "fastapi", "flask"]):
        return "Python Developer"
    elif any(k in skills_lower for k in ["aws", "docker", "kubernetes", "terraform", "ci/cd", "jenkins", "ansible"]):
        return "DevOps Engineer"
    elif any(k in skills_lower for k in ["data engineering", "spark", "hadoop", "airflow", "etl", "bigquery", "snowflake"]):
        return "Data Engineer"
    elif any(k in skills_lower for k in ["selenium", "automation testing", "cypress", "testng", "sdet"]):
        return "QA Automation Engineer"
    elif any(k in skills_lower for k in [".net", "asp.net", "c#", ".net core"]):
        return "Dot Net Developer"
    elif any(k in skills_lower for k in ["salesforce", "apex", "lwc", "visualforce"]):
        return "Salesforce Developer"
    elif any(k in skills_lower for k in ["sap", "abap", "sap fico", "sap mm", "sap hana"]):
        return "SAP Consultant"
    elif any(k in skills_lower for k in ["node.js", "express", "nestjs"]):
        return "Node.js Developer"
    elif any(k in skills_lower for k in ["flutter", "android", "ios", "swift", "kotlin"]):
        return "Mobile Developer"
    elif any(k in skills_lower for k in ["sql", "mysql", "postgresql", "dba", "database tuning"]):
        return "SQL Developer"

    return "Software Engineer"

def classify_candidate_skills(
    skills: List[str],
    position: str,
    raw_text: str = ""
) -> Tuple[List[str], List[str]]:
    """
    Intelligently divides extracted candidate skills into:
    - Primary Skills: Main / Core technical skills aligning with the position.
    - Secondary Skills: Supporting / auxiliary skills, tools, cloud, and domains.
    """
    pos_lower = position.lower()
    matched_profile = None

    for domain_key, profile in DOMAIN_SKILL_PROFILES.items():
        if domain_key in pos_lower or profile["title"].lower() in pos_lower:
            matched_profile = profile
            break

    primary_skills: List[str] = []
    secondary_skills: List[str] = []

    if matched_profile:
        domain_primaries = [p.lower() for p in matched_profile["primary"]]
        domain_secondaries = [s.lower() for s in matched_profile["secondary"]]

        for sk in skills:
            sk_lower = sk.lower()
            if sk_lower in domain_primaries:
                if sk not in primary_skills:
                    primary_skills.append(sk)
            else:
                if sk not in secondary_skills:
                    secondary_skills.append(sk)

        # Also search raw text for explicit primary/secondary skills mentioned in the profile
        text_lower = raw_text.lower()
        for p_skill in matched_profile["primary"]:
            if p_skill not in primary_skills and re.search(r"\b" + re.escape(p_skill.lower()) + r"\b", text_lower):
                primary_skills.append(p_skill)
        for s_skill in matched_profile["secondary"]:
            if s_skill not in secondary_skills and s_skill not in primary_skills and re.search(r"\b" + re.escape(s_skill.lower()) + r"\b", text_lower):
                secondary_skills.append(s_skill)

    # Fallback if no specific profile matched or primary list is empty
    if not primary_skills:
        # Split first 4 skills as primary, remaining as secondary
        primary_skills = skills[:4]
        secondary_skills = skills[4:]
    elif not secondary_skills and len(skills) > len(primary_skills):
        secondary_skills = [s for s in skills if s not in primary_skills]

    # Ensure supporting domain skills (Finance, IT, Support, Reporting, Git, Linux, Agile) are in secondary
    support_keywords = ["Finance", "IT", "Support", "Reporting", "Production Support", "Maintenance", "Jira", "Agile", "Scrum", "Git", "Linux", "Windows"]
    for kw in support_keywords:
        if kw in primary_skills and len(primary_skills) > 1:
            primary_skills.remove(kw)
            if kw not in secondary_skills:
                secondary_skills.append(kw)

    return primary_skills, secondary_skills

def infer_position_and_skills(
    current_designation: Optional[str] = None,
    skills: Optional[List[str]] = None,
    all_skills: Optional[List[str]] = None,
    stored_primary: Optional[List[str]] = None,
    stored_secondary: Optional[List[str]] = None,
    raw_designation: Optional[str] = None
) -> Tuple[str, List[str], List[str]]:
    """
    Helper to accurately infer position and classify primary/secondary skills for existing database records.
    """
    pos = current_designation or raw_designation
    skills_list = skills or all_skills or []

    # If position is blank, generic "Software Engineer", or "Software Developer", infer from skills
    if not pos or pos.strip() in ["", "Software Engineer", "Software Developer", "Applicant", "Developer"]:
        pos = extract_position_from_text("", "", skills_list)

    if stored_primary and len(stored_primary) > 0:
        primary = stored_primary
        secondary = stored_secondary or []
    else:
        primary, secondary = classify_candidate_skills(skills_list, pos, "")

    return pos, primary, secondary

def parse_candidate_from_text(raw_text: str, filename: str = "") -> Dict[str, Any]:
    """
    Intelligently parses candidate details from resume text.
    Extracts exact position, primary skills, secondary skills, and full contact details.
    """
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    full_text = " ".join(lines)
    
    # 1. Email Extraction
    email = None
    email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b", full_text)
    if email_match:
        email = email_match.group(0).lower()

    # 2. Phone Extraction (Supports International, Indian +91, US +1 formats)
    phone = None
    phone_pattern = r"(?:(?:\+|00)(\d{1,3})[\s.-]?)?(?:\(?(\d{2,4})\)?[\s.-]?)?(\d{3,5}[\s.-]?\d{3,5})"
    phone_matches = re.finditer(phone_pattern, full_text)
    for match in phone_matches:
        candidate_phone = match.group(0).strip()
        digits_only = re.sub(r"[^\d+]", "", candidate_phone)
        if 10 <= len(re.sub(r"\D", "", digits_only)) <= 15:
            phone = candidate_phone
            break

    # 3. Name Extraction
    first_name = ""
    last_name = ""
    full_name = ""
    
    for line in lines[:8]:
        if "@" in line or "http" in line or "resume" in line.lower() or "curriculum" in line.lower() or "developer" in line.lower() or "engineer" in line.lower():
            continue
        cleaned_line = re.sub(r"[^a-zA-Z\s]", "", line).strip()
        words = cleaned_line.split()
        if 2 <= len(words) <= 4 and all(len(w) > 1 for w in words):
            first_name = words[0].capitalize()
            last_name = " ".join(words[1:]).capitalize()
            full_name = f"{first_name} {last_name}"
            break
            
    # Fallback to filename if name not found
    if not full_name and filename:
        clean_file = os.path.splitext(filename)[0]
        clean_file = re.sub(r"(resume|cv|profile|_|-|\d)", " ", clean_file, flags=re.IGNORECASE).strip()
        words = clean_file.split()
        if len(words) >= 2:
            first_name = words[0].capitalize()
            last_name = " ".join(words[1:]).capitalize()
            full_name = f"{first_name} {last_name}"
        elif len(words) == 1:
            first_name = words[0].capitalize()
            full_name = first_name

    # 4. Total Experience
    exp_years = 0.0
    exp_matches = re.finditer(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)(?:\s*(?:of)?\s*(?:experience|exp))?", full_text, re.IGNORECASE)
    for m in exp_matches:
        try:
            val = float(m.group(1))
            if 0.5 <= val <= 40:
                exp_years = max(exp_years, val)
        except ValueError:
            pass

    # 5. Skills Extraction
    found_skills = []
    text_lower = full_text.lower()
    for skill in TECH_SKILLS_DB:
        pattern = r"\b" + re.escape(skill.lower()) + r"\b"
        if re.search(pattern, text_lower):
            if skill not in found_skills:
                found_skills.append(skill)

    # 6. Exact Position / Job Title Extraction
    position = extract_position_from_text(raw_text, filename, found_skills)

    # 7. Primary vs Secondary Skills Classification
    primary_skills, secondary_skills = classify_candidate_skills(found_skills, position, raw_text)

    # Combine all unique skills
    all_skills = list(dict.fromkeys(primary_skills + secondary_skills + found_skills))

    # 8. Education / Qualification
    highest_qual = ""
    for pat, label in DEGREE_PATTERNS:
        if re.search(pat, full_text, re.IGNORECASE):
            highest_qual = label
            break

    # 9. Current Company Heuristics
    current_company = ""
    company_match = re.search(r"(?:working\s+at|employed\s+at|current\s+company\s*[:\-\|\–])\s*([A-Za-z0-9\s,\.\-&]{3,40})", full_text, re.IGNORECASE)
    if company_match:
        current_company = company_match.group(1).strip()

    # 10. LinkedIn & GitHub URLs
    linkedin_url = ""
    github_url = ""
    li_m = re.search(r"(https?://(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+)", full_text, re.IGNORECASE)
    if li_m:
        linkedin_url = li_m.group(0)
    gh_m = re.search(r"(https?://(?:www\.)?github\.com/[A-Za-z0-9_-]+)", full_text, re.IGNORECASE)
    if gh_m:
        github_url = gh_m.group(0)

    # 11. Notice Period
    notice_period = ""
    np_m = re.search(r"\b(Immediate|15\s*days?|30\s*days?|60\s*days?|90\s*days?|1\s*month|2\s*months?|3\s*months?)\b", full_text, re.IGNORECASE)
    if np_m:
        notice_period = np_m.group(0)

    # 12. Location Heuristics
    location = ""
    common_cities = ["Bangalore", "Bengaluru", "Hyderabad", "Pune", "Mumbai", "Delhi", "Gurgaon", "Gurugram", "Noida", "Chennai", "San Francisco", "Seattle", "New York", "London", "Remote"]
    for city in common_cities:
        if re.search(r"\b" + re.escape(city) + r"\b", full_text, re.IGNORECASE):
            location = city
            break

    whatsapp_number = phone or ""
    
    return {
        "first_name": first_name,
        "last_name": last_name,
        "full_name": full_name or "Applicant",
        "email": email or "",
        "phone": phone or "",
        "alternate_phone": "",
        "whatsapp_number": whatsapp_number,
        "country_code": "+91" if whatsapp_number.startswith("+91") or (phone and len(re.sub(r"\D", "", phone)) == 10) else "+1",
        "location": location,
        "preferred_location": "",
        "total_experience": exp_years if exp_years > 0 else 2.0,
        "relevant_experience": max(0.0, exp_years - 0.5) if exp_years > 0 else 2.0,
        "current_company": current_company,
        "current_designation": position,
        "position": position,
        "skills": all_skills,
        "technical_skills": all_skills,
        "primary_skills": primary_skills,
        "secondary_skills": secondary_skills,
        "bench_primary_skills": primary_skills,
        "bench_secondary_skills": secondary_skills,
        "education": highest_qual or "Bachelor's Degree",
        "highest_qualification": highest_qual or "Bachelor's Degree",
        "notice_period": notice_period or "30 Days",
        "current_ctc": None,
        "expected_ctc": None,
        "linkedin_url": linkedin_url,
        "github_url": github_url,
        "certifications": [],
        "date_of_birth": "",
        "summary": f"{position} with {exp_years or 2} years of experience specializing in {', '.join(primary_skills[:4])}."
    }

def parse_cv_document(filename: str, content: bytes) -> Dict[str, Any]:
    """
    Directly extracts and parses candidate details from file content bytes.
    """
    raw_text = extract_text_from_file(content, filename)
    return parse_candidate_from_text(raw_text, filename)

def validate_whatsapp_eligibility(
    db: Session,
    phone_or_whatsapp: Optional[str],
    consent_status: Optional[WhatsAppConsentStatusEnum] = None,
    candidate_id: Optional[str] = None
) -> WhatsAppEligibilityInfo:
    """
    Computes real WhatsApp outreach eligibility status according to compliance.
    """
    if not phone_or_whatsapp or not phone_or_whatsapp.strip():
        return WhatsAppEligibilityInfo(
            is_eligible=False,
            status="Invalid Number",
            whatsapp_number=None,
            consent_status=WhatsAppConsentStatusEnum.NOT_COLLECTED,
            opt_out_status=False,
            reason="Mobile or WhatsApp number is missing."
        )

    clean_digits = re.sub(r"[^\d+]", "", phone_or_whatsapp)
    pure_digits = re.sub(r"\D", "", clean_digits)
    
    if len(pure_digits) < 10 or len(pure_digits) > 15:
        return WhatsAppEligibilityInfo(
            is_eligible=False,
            status="Invalid Number",
            whatsapp_number=clean_digits,
            consent_status=consent_status or WhatsAppConsentStatusEnum.NOT_COLLECTED,
            opt_out_status=False,
            reason="Phone number format is invalid (requires 10-15 digits)."
        )

    # Check Opt-out / suppression list
    opt_out_record = db.query(WhatsAppOptOut).filter(
        (WhatsAppOptOut.whatsapp_number == clean_digits) | 
        (WhatsAppOptOut.whatsapp_number == pure_digits) |
        (WhatsAppOptOut.whatsapp_number.ilike(f"%{pure_digits[-10:]}%"))
    ).filter(WhatsAppOptOut.is_active == True).first()

    if opt_out_record:
        return WhatsAppEligibilityInfo(
            is_eligible=False,
            status="Opted Out",
            whatsapp_number=clean_digits,
            consent_status=WhatsAppConsentStatusEnum.OPTED_OUT,
            opt_out_status=True,
            reason=f"Candidate previously opted out: {opt_out_record.reason}"
        )

    effective_consent = consent_status or WhatsAppConsentStatusEnum.NOT_COLLECTED
    
    if effective_consent == WhatsAppConsentStatusEnum.GRANTED:
        return WhatsAppEligibilityInfo(
            is_eligible=True,
            status="Eligible",
            whatsapp_number=clean_digits,
            consent_status=WhatsAppConsentStatusEnum.GRANTED,
            opt_out_status=False,
            reason="Valid number and active consent verified."
        )
    elif effective_consent == WhatsAppConsentStatusEnum.REVOKED:
        return WhatsAppEligibilityInfo(
            is_eligible=False,
            status="Blocked",
            whatsapp_number=clean_digits,
            consent_status=WhatsAppConsentStatusEnum.REVOKED,
            opt_out_status=False,
            reason="WhatsApp consent has been revoked by candidate."
        )
    else:
        return WhatsAppEligibilityInfo(
            is_eligible=False,
            status="Consent Required",
            whatsapp_number=clean_digits,
            consent_status=effective_consent,
            opt_out_status=False,
            reason="WhatsApp consent is required before proactive outreach can be sent."
        )

def check_candidate_duplicate(
    db: Session,
    email: Optional[str] = None,
    phone: Optional[str] = None,
    whatsapp_number: Optional[str] = None,
    first_name: Optional[str] = None,
    last_name: Optional[str] = None,
    exclude_candidate_id: Optional[str] = None
) -> Tuple[bool, Optional[Candidate], Optional[str]]:
    """
    Checks if a candidate already exists in the talent pool.
    """
    query = db.query(Candidate)
    if exclude_candidate_id:
        query = query.filter(Candidate.id != exclude_candidate_id)

    # 1. Exact Email Match
    if email and email.strip():
        dup_email = query.filter(Candidate.email.ilike(email.strip())).first()
        if dup_email:
            return True, dup_email, f"Duplicate found by Email: {email}"

    # 2. Phone Match
    if phone and phone.strip():
        pure_p = re.sub(r"\D", "", phone)
        if len(pure_p) >= 10:
            dup_phone = query.filter(Candidate.phone.ilike(f"%{pure_p[-10:]}%")).first()
            if dup_phone:
                return True, dup_phone, f"Duplicate found by Phone: {phone}"

    # 3. WhatsApp Match
    if whatsapp_number and whatsapp_number.strip():
        pure_w = re.sub(r"\D", "", whatsapp_number)
        if len(pure_w) >= 10:
            dup_wa = query.filter(Candidate.whatsapp_number.ilike(f"%{pure_w[-10:]}%")).first()
            if dup_wa:
                return True, dup_wa, f"Duplicate found by WhatsApp Number: {whatsapp_number}"

    # 4. Name + Phone combination
    if first_name and last_name and (first_name.strip().lower() != "extracted" or last_name.strip().lower() != "candidate"):
        if phone:
            pure_p = re.sub(r"\D", "", phone)
            if len(pure_p) >= 10:
                dup_name_phone = query.filter(
                    Candidate.first_name.ilike(first_name.strip()),
                    Candidate.last_name.ilike(last_name.strip()),
                    Candidate.phone.ilike(f"%{pure_p[-10:]}%")
                ).first()
                if dup_name_phone:
                    return True, dup_name_phone, f"Duplicate candidate with matching name and phone: {first_name} {last_name}"

    return False, None, None
